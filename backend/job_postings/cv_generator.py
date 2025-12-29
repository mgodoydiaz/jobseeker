from docx import Document
from typing import Dict

def generate_cv_from_template(template_path: str, context: Dict[str, str], output_path: str):
    """
    Genera un currículum en formato .docx a partir de una plantilla y un contexto de datos.

    Busca y reemplaza placeholders en párrafos y tablas. 
    Los placeholders deben tener el formato {{KEY}}.

    Args:
        template_path: La ruta al archivo de plantilla .docx.
        context: Un diccionario donde las claves son los placeholders (sin los corchetes)
                 y los valores son los datos con los que se reemplazarán.
        output_path: La ruta donde se guardará el documento .docx generado.
    """
    try:
        # Abrir el documento de plantilla
        doc = Document(template_path)

        # Reemplazar placeholders en los párrafos
        for paragraph in doc.paragraphs:
            for key, value in context.items():
                # Construir el formato del placeholder (ej: {{NOMBRE}})
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # Usamos `inline.runs` para preservar el formato del texto
                    for run in paragraph.runs:
                        run.text = run.text.replace(placeholder, value)

        # Reemplazar placeholders en las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in context.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Guardar el documento generado
        doc.save(output_path)
        print(f"Currículum generado y guardado en: {output_path}")
        return True

    except Exception as e:
        print(f"Error al generar el currículum: {e}")
        return False


# Ejemplo de uso (esto no se ejecutará cuando se importe como módulo)
if __name__ == '__main__':
    # Crear un contexto de ejemplo con los datos a rellenar
    cv_context = {
        "NOMBRE": "Maria Garcia",
        "TELEFONO": "+34 600 123 456",
        "EMAIL": "maria.garcia@email.com",
        "LINKEDIN": "linkedin.com/in/mariagarcia",
        "RESUMEN": "Desarrolladora de software proactiva con 3 años de experiencia en el desarrollo de aplicaciones web con Python y Django. Busco oportunidades para aplicar mis habilidades en un entorno desafiante.",
        "PUESTO_1": "Desarrolladora de Software Junior",
        "EMPRESA_1": "Tech Solutions S.L.",
        "FECHA_1": "Enero 2021 - Presente",
        "DESC_1": "- Desarrollo y mantenimiento de funcionalidades para una plataforma de e-commerce.\n- Colaboración en el diseño de la arquitectura de la API RESTful.",
        "PUESTO_2": "Becaria de Desarrollo",
        "EMPRESA_2": "Innovatech",
        "FECHA_2": "Junio 2020 - Diciembre 2020",
        "DESC_2": "- Apoyo en la refactorización de código y corrección de bugs.",
        "TITULACION_1": "Grado en Ingeniería Informática",
        "UNIVERSIDAD_1": "Universidad Politécnica de Madrid",
        "FECHA_UNI_1": "2016 - 2020",
        "IDIOMA_1": "Inglés (Nivel C1)",
        "IDIOMA_2": "Alemán (Nivel A2)",
    }

    # Rutas a los archivos (suponiendo que la plantilla está en la raíz del proyecto)
    template_file = "cv_template.docx" 
    output_file = "cv_generado_maria_garcia.docx"

    # Generar el CV
    # Nota: Este script espera que exista un archivo `cv_template.docx` con los placeholders correspondientes.
    # Como no podemos subir archivos aquí, este script fallará si se ejecuta directamente.
    print("Intentando generar CV a partir de la plantilla...")
    print("Este script de ejemplo necesita un archivo `cv_template.docx` para funcionar.")
    # generate_cv_from_template(template_file, cv_context, output_file)

